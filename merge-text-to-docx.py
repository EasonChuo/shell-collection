import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import chardet
import time

root_path = 'C:/Users/chuon/Downloads/delphiCode'
docx_name = 'output.docx'
ext = ['.ddp',
    '.dti',
    '.pas',
    '.dof',
    '.dpr',
    '.dof',
    '.cfg'
]

def find_files(directory, extensions):
    """Recursively find files with given extensions in directory."""
    files_found = []
    for root, dirs, files in os.walk(directory):
        for file in files:
            lowercase_filename = file.lower()
            if any(lowercase_filename.endswith(ext) for ext in extensions):
                files_found.append(os.path.join(root, file))
    return files_found

# 加入目錄
def add_toc(document):
    """Add a table of contents (TOC) to the document."""
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldSimple')
    fldChar.set(qn('w:instr'), 'TOC \\o "1-3" \\h \\z \\u')
    run._r.append(fldChar)
    document.add_page_break()

def add_footer_with_page_numbers(section):
    """Add a footer with page numbers to a section."""
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.text = "Page "
    run = paragraph.add_run()
    field_code = 'PAGE'
    docx_field = OxmlElement('w:fldSimple')
    docx_field.set(qn('w:instr'), field_code)
    run._r.append(docx_field)

def clean_text(text):
    # 移除NULL字節和控制字符，只保留ASCII字符
    return ''.join(char for char in text if char.isprintable() or char.isspace())

def main(directory, extensions, output_file):
    document = Document()
    add_toc(document)  # Add TOC at the beginning
    files_found = find_files(directory, extensions)
    total_lines = 0
    # 計算檔案處理數量
    total_files = len(files_found)
    file_counter = 0

    for file_path in files_found:
        # 計算檔案處理數量
        file_counter = file_counter + 1
        print(f'{file_counter}/{total_files} {file_path}')

        #自動偵測檔案編碼
        with open(file_path, 'rb') as file:
            raw_data = file.read()
            encoding = chardet.detect(raw_data)['encoding']

        # print(f'{encoding} {file_path}')
        with open(file_path, 'r', encoding=encoding, errors='replace') as file:
            contents = file.readlines()
            line_count = len(contents)
            total_lines += line_count
            section = document.sections[-1]
            add_footer_with_page_numbers(section)
            
            # Add file name as a new section heading for TOC
            heading = document.add_heading(level=1)
            # heading_run = heading.add_run(os.path.basename(file_path))
            heading_run = heading.add_run(file_path.replace(root_path, ''))
            heading_run.bold = True
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            document.add_paragraph(f'Line Count: {line_count}')
            for line in contents:
                # document.add_paragraph(line, style='BodyText')
                document.add_paragraph(clean_text(line).splitlines(), style='BodyText')
                # document.add_paragraph("".join(line), style='BodyText')
            document.add_page_break()

    # Add a placeholder for total line count at the end
    document.add_paragraph(f'Total lines across all files: {total_lines}')
    print(f'Total lines across all files: {total_lines}')

    document.save(output_file)
    print(f'Total lines across all files: {total_lines}')

if __name__ == "__main__":
    start_time = time.time()  # 程式開始前的時間
    main(root_path, ext, docx_name)
    end_time = time.time()  # 程式結束的時間
    elapsed_time = end_time - start_time  # 計算耗時
    print(f"程式執行耗時: {elapsed_time} 秒")